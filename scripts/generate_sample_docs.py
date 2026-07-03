"""Generates sample TOR and Agreement Template .docx fixtures for manual testing.

Run with:
    python scripts/generate_sample_docs.py

Writes `samples/TOR_Sample.docx` and `samples/Agreement_Template_Sample.docx`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"


def build_tor() -> Document:
    doc = Document()
    doc.add_heading("Terms of Reference — Beneficiary Data Systems Consultancy", level=0)

    doc.add_heading("Background", level=1)
    doc.add_paragraph(
        "The organization is implementing a multi-district program to strengthen "
        "beneficiary data collection and reporting. This consultancy will support "
        "the design and rollout of a digital data collection system."
    )

    doc.add_heading("Scope of Work", level=1)
    doc.add_paragraph(
        "The consultant will design, build, and roll out a mobile-based beneficiary "
        "data collection system across 5 program districts, and will train field "
        "staff on its use."
    )
    doc.add_paragraph(
        "The consultant will also conduct a data quality assessment of existing "
        "records and recommend improvements to data collection protocols."
    )

    doc.add_heading("Deliverables", level=1)
    doc.add_paragraph("Data collection system design document", style="List Number")
    doc.add_paragraph("Fully functional mobile data collection application", style="List Number")
    doc.add_paragraph("Field staff training curriculum and completion report", style="List Number")
    doc.add_paragraph("Data quality assessment report with recommendations", style="List Number")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Deliverable"
    header_cells[1].text = "Due Date"
    for deliverable, due in [
        ("Design document", "Week 4"),
        ("Mobile application (v1)", "Week 10"),
        ("Training completion report", "Week 14"),
        ("Data quality assessment report", "Week 16"),
    ]:
        row_cells = table.add_row().cells
        row_cells[0].text = deliverable
        row_cells[1].text = due

    doc.add_heading("Timeline", level=1)
    doc.add_paragraph("The assignment shall be completed within 16 weeks of the contract start date.")

    doc.add_heading("Payment Schedule", level=1)
    doc.add_paragraph(
        "Payments will be made in three installments: 30% upon signing, 40% upon "
        "delivery of the mobile application, and 30% upon submission of the final "
        "training and assessment reports."
    )

    doc.add_heading("Reporting Requirements", level=1)
    doc.add_paragraph(
        "The consultant shall submit a brief progress report every two weeks and a "
        "final consolidated report at the end of the assignment."
    )

    doc.add_heading("Qualifications Required", level=1)
    doc.add_paragraph(
        "At least 5 years of experience building mobile data collection tools for "
        "development-sector programs, with demonstrated experience in similar "
        "multi-district rollouts."
    )

    return doc


def build_agreement_template() -> Document:
    doc = Document()
    doc.add_heading("Consultancy Agreement", level=0)
    doc.add_paragraph(
        "This Agreement is made between [Organization Name] (\"the Organization\") and "
        "[Consultant Name] (\"the Consultant\")."
    )

    doc.add_heading("Scope of Services", level=1)
    doc.add_paragraph("[Placeholder — to be replaced with the scope of services for this engagement.]")

    doc.add_heading("Deliverables & Timelines", level=1)
    doc.add_paragraph("[Placeholder — to be replaced with the agreed deliverables and timelines.]")

    doc.add_heading("Payment Terms", level=1)
    doc.add_paragraph("[Placeholder — to be replaced with the agreed payment terms.]")

    doc.add_heading("Reporting Obligations", level=1)
    doc.add_paragraph("[Placeholder — to be replaced with the agreed reporting obligations.]")

    doc.add_heading("Confidentiality", level=1)
    doc.add_paragraph(
        "The Consultant shall keep confidential all non-public information obtained "
        "during the course of this engagement and shall not disclose it to any third "
        "party without prior written consent."
    )

    doc.add_heading("Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this Agreement by providing 30 days' written "
        "notice to the other party."
    )

    doc.add_heading("Governing Law and Dispute Resolution", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by the laws of India. Any disputes shall "
        "be resolved through arbitration in accordance with the Arbitration and "
        "Conciliation Act, 1996."
    )

    return doc


def main() -> None:
    SAMPLES_DIR.mkdir(exist_ok=True)
    tor_path = SAMPLES_DIR / "TOR_Sample.docx"
    template_path = SAMPLES_DIR / "Agreement_Template_Sample.docx"

    build_tor().save(tor_path)
    build_agreement_template().save(template_path)

    print(f"Wrote {tor_path}")
    print(f"Wrote {template_path}")


if __name__ == "__main__":
    main()
